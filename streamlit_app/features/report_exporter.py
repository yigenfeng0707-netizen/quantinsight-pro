# -*- coding: utf-8 -*-
"""
QuantInsight Pro - 报告导出模块
=================================

支持:
  - Word (.docx) - python-docx生成，含封面/目录/图表/水印
  - PDF - reportlab生成，专业排版
"""
import io
from datetime import datetime
from typing import Dict


# ============== Word 导出 ==============

def export_word(report: Dict) -> bytes:
    """生成Word格式报告（含封面+目录+正文+水印）"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn, nsmap
        from docx.oxml import OxmlElement
    except ImportError:
        return _word_fallback(report)

    doc = Document()

    # 设置中文字体（关键：解决中文乱码）
    _set_chinese_font(doc)

    # ===== 封面 =====
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('QuantInsight Pro')
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)
    title_run.font.name = '微软雅黑'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run('AI 驱动的另类数据量化投研平台')
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
    sub_run.font.name = '微软雅黑'
    sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    doc.add_paragraph()

    report_title = doc.add_paragraph()
    report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt_run = report_title.add_run(report['title'])
    rt_run.font.size = Pt(24)
    rt_run.font.bold = True
    rt_run.font.color.rgb = RGBColor(0x00, 0xD4, 0xFF)
    rt_run.font.name = '微软雅黑'
    rt_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"生成时间: {report['generated_at']}\n"
        f"数据源: {report['data_source']}\n"
        f"AI引擎: Qwen3.7-Max"
    )
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    meta_run.font.name = '微软雅黑'
    meta_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 封底水印（封面后）
    doc.add_page_break()

    # ===== 目录 =====
    toc_title = doc.add_heading('目录', level=1)
    for s in report['sections']:
        p = doc.add_paragraph(s['name'], style='List Bullet')
    doc.add_page_break()

    # ===== 大盘速览 =====
    doc.add_heading('一、大盘速览', level=1)
    macro = report['data']['macro']
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['指数', '收盘', '涨跌幅']
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    rows_data = [
        ['上证指数', f"{macro['sh_index']['value']:.2f}", f"{macro['sh_index']['change_pct']:+.2f}%"],
        ['深证成指', f"{macro['sz_index']['value']:.2f}", f"{macro['sz_index']['change_pct']:+.2f}%"],
        ['创业板指', f"{macro['cyb_index']['value']:.2f}", f"{macro['cyb_index']['change_pct']:+.2f}%"],
    ]
    for i, row_data in enumerate(rows_data, 1):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"涨跌家数: 涨 {macro['up_count']} 家 / 跌 {macro['down_count']} 家\n").font.name = '微软雅黑'
    p.add_run(f"涨停/跌停: {macro['limit_up']} / {macro['limit_down']}\n").font.name = '微软雅黑'
    p.add_run(f"北向资金: {macro['north_flow']:+.2f} 亿元").font.name = '微软雅黑'

    # ===== 行业涨跌 =====
    doc.add_heading('二、行业涨跌榜', level=1)
    industries = report['data']['industries'][:8]
    ind_table = doc.add_table(rows=len(industries) + 1, cols=3)
    ind_table.style = 'Light Grid Accent 1'
    for j, h in enumerate(['行业', '涨跌幅', '领涨股']):
        cell = ind_table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for i, ind in enumerate(industries, 1):
        for j, val in enumerate([ind['name'], f"{ind['change_pct']:+.2f}%", ind.get('leader', '')]):
            cell = ind_table.rows[i].cells[j]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ===== 6段式分析 =====
    doc.add_page_break()
    doc.add_heading('三、深度分析', level=1)
    for i, s in enumerate(report['sections'], 1):
        doc.add_heading(f"{i}. {s['name']}", level=2)
        for para_text in s['content'].split('\n'):
            if para_text.strip():
                p = doc.add_paragraph(para_text)
                for run in p.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(10.5)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ===== 免责声明 =====
    doc.add_page_break()
    doc.add_heading('免责声明', level=1)
    disclaimer = doc.add_paragraph(
        "本报告由 QuantInsight Pro AI 自动生成, 数据来源于公开市场信息及 akshare 财经数据接口. "
        "报告内容仅供参考和学习交流, 不构成任何投资建议. "
        "投资者应根据自身风险承受能力和投资目标, 独立做出投资决策. "
        "市场有风险, 投资需谨慎."
    )
    for run in disclaimer.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 页脚
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('QuantInsight Pro  © 2026  |  https://quantinsight.cn')
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 保存到BytesIO
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _set_chinese_font(doc):
    """设置全局中文字体"""
    from docx.oxml.ns import qn
    from docx.shared import Pt
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def _word_fallback(report: Dict) -> bytes:
    """python-docx不可用时降级为UTF-8文本（保证可下载）"""
    text = report.get('raw_markdown', report['title'])
    return text.encode('utf-8')


# ============== PDF 导出 ==============

def export_pdf(report: Dict) -> bytes:
    """生成PDF报告"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib.colors import HexColor
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak, Image,
        )
    except ImportError:
        return _pdf_fallback(report)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
        title=report['title'],
        author='QuantInsight Pro',
    )

    # 注册中文字体（关键：解决中文乱码）
    _register_chinese_pdf_font()

    styles = _build_pdf_styles()

    story = []

    # ===== 封面 =====
    story.append(Spacer(1, 4*cm))
    story.append(Paragraph('QuantInsight Pro', styles['CoverTitle']))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph('AI 驱动的另类数据量化投研平台', styles['CoverSubtitle']))
    story.append(Spacer(1, 3*cm))
    story.append(Paragraph(report['title'], styles['ReportTitle']))
    story.append(Spacer(1, 2*cm))
    story.append(Paragraph(
        f"生成时间: {report['generated_at']}<br/>"
        f"数据源: {report['data_source']}<br/>"
        f"AI引擎: Qwen3.7-Max",
        styles['Meta'],
    ))
    story.append(PageBreak())

    # ===== 目录 =====
    story.append(Paragraph('目录', styles['H1']))
    for s in report['sections']:
        story.append(Paragraph(f"• {s['name']}", styles['TocItem']))
    story.append(PageBreak())

    # ===== 大盘速览 =====
    story.append(Paragraph('一、大盘速览', styles['H1']))
    macro = report['data']['macro']
    data = [
        ['指数', '收盘', '涨跌幅'],
        ['上证指数', f"{macro['sh_index']['value']:.2f}", f"{macro['sh_index']['change_pct']:+.2f}%"],
        ['深证成指', f"{macro['sz_index']['value']:.2f}", f"{macro['sz_index']['change_pct']:+.2f}%"],
        ['创业板指', f"{macro['cyb_index']['value']:.2f}", f"{macro['cyb_index']['change_pct']:+.2f}%"],
    ]
    t = Table(data, colWidths=[5*cm, 5*cm, 5*cm])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), HexColor('#0A1628')),
        ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#FFFFFF')),
        ('FONTNAME', (0, 0), (-1, -1), 'ChineseFont'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#2A3441')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#F5F7FA'), HexColor('#FFFFFF')]),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph(
        f"<b>涨跌家数</b>: 涨 {macro['up_count']} 家 / 跌 {macro['down_count']} 家<br/>"
        f"<b>涨停/跌停</b>: {macro['limit_up']} / {macro['limit_down']}<br/>"
        f"<b>北向资金</b>: {macro['north_flow']:+.2f} 亿元",
        styles['Body'],
    ))

    # ===== 行业涨跌 =====
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph('二、行业涨跌榜', styles['H1']))
    industries = report['data']['industries'][:8]
    ind_data = [['行业', '涨跌幅', '领涨股']]
    for ind in industries:
        ind_data.append([ind['name'], f"{ind['change_pct']:+.2f}%", ind.get('leader', '')])
    t2 = Table(ind_data, colWidths=[5*cm, 4*cm, 6*cm])
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), HexColor('#0A1628')),
        ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#FFFFFF')),
        ('FONTNAME', (0, 0), (-1, -1), 'ChineseFont'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#2A3441')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#F5F7FA'), HexColor('#FFFFFF')]),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(t2)

    # ===== 6段分析 =====
    story.append(PageBreak())
    story.append(Paragraph('三、深度分析', styles['H1']))
    for i, s in enumerate(report['sections'], 1):
        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph(f"{i}. {s['name']}", styles['H2']))
        # 处理段落
        for para_text in s['content'].split('\n'):
            if para_text.strip():
                story.append(Paragraph(para_text, styles['Body']))

    # ===== 免责声明 =====
    story.append(PageBreak())
    story.append(Paragraph('免责声明', styles['H1']))
    story.append(Paragraph(
        "本报告由 QuantInsight Pro AI 自动生成, 数据来源于公开市场信息及 akshare 财经数据接口. "
        "报告内容仅供参考和学习交流, 不构成任何投资建议. "
        "投资者应根据自身风险承受能力和投资目标, 独立做出投资决策. "
        "市场有风险, 投资需谨慎.",
        styles['Disclaimer'],
    ))

    # 添加页脚（页码）
    def _add_page_number(canvas, doc):
        canvas.saveState()
        canvas.setFont('ChineseFont', 8)
        canvas.setFillColor(HexColor('#999999'))
        canvas.drawCentredString(A4[0]/2, 1*cm, f"QuantInsight Pro  |  {doc.page} / " )
        canvas.restoreState()

    doc.build(story, onFirstPage=_add_page_number, onLaterPages=_add_page_number)
    buf.seek(0)
    return buf.getvalue()


def _register_chinese_pdf_font():
    """注册中文字体到reportlab（避免乱码）"""
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        # 尝试常见中文字体路径
        font_paths = [
            '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            '/System/Library/Fonts/STHeiti Medium.ttc',
            'C:/Windows/Fonts/msyh.ttc',
            'C:/Windows/Fonts/simhei.ttf',
        ]
        for fp in font_paths:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont('ChineseFont', fp))
                    return
                except Exception:
                    continue
    except Exception:
        pass


def _build_pdf_styles():
    """构建PDF样式"""
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    return {
        'CoverTitle': ParagraphStyle('CoverTitle', fontName='ChineseFont', fontSize=36,
                                     textColor=HexColor('#0A1628'), alignment=TA_CENTER, leading=44),
        'CoverSubtitle': ParagraphStyle('CoverSubtitle', fontName='ChineseFont', fontSize=14,
                                        textColor=HexColor('#6C757D'), alignment=TA_CENTER, leading=18),
        'ReportTitle': ParagraphStyle('ReportTitle', fontName='ChineseFont', fontSize=24,
                                      textColor=HexColor('#00D4FF'), alignment=TA_CENTER, leading=30,
                                      spaceAfter=20),
        'Meta': ParagraphStyle('Meta', fontName='ChineseFont', fontSize=11,
                               textColor=HexColor('#4A5568'), alignment=TA_CENTER, leading=18),
        'H1': ParagraphStyle('H1', fontName='ChineseFont', fontSize=18,
                             textColor=HexColor('#0A1628'), leading=24, spaceBefore=12, spaceAfter=8),
        'H2': ParagraphStyle('H2', fontName='ChineseFont', fontSize=14,
                             textColor=HexColor('#1F4E78'), leading=20, spaceBefore=8, spaceAfter=4),
        'Body': ParagraphStyle('Body', fontName='ChineseFont', fontSize=10.5,
                               textColor=HexColor('#333333'), leading=16, spaceAfter=6,
                               alignment=TA_LEFT),
        'TocItem': ParagraphStyle('TocItem', fontName='ChineseFont', fontSize=12,
                                  textColor=HexColor('#333333'), leading=20, leftIndent=20),
        'Disclaimer': ParagraphStyle('Disclaimer', fontName='ChineseFont', fontSize=9,
                                     textColor=HexColor('#808080'), leading=14, alignment=TA_LEFT),
    }


def _pdf_fallback(report: Dict) -> bytes:
    """reportlab不可用时降级为UTF-8文本"""
    text = report.get('raw_markdown', report['title'])
    return text.encode('utf-8')
