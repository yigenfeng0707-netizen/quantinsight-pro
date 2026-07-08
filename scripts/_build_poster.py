"""
P2 锦上添花物料生成（A1 海报 + 易拉宝 + 永字资管背书牌）
- 海报 PNG: 1190x1684 px @ 50dpi  (23.8 x 33.68 inch)
- 易拉宝 PNG: 1000x2500 px @ 50dpi  (20 x 50 inch)
- 永字资管背书函 DOCX
"""
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, Rectangle, Polygon
import matplotlib.font_manager as fm
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 中文字体
plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

# 配色
PRIMARY = '#1F77B4'
DARK = '#1A1A1A'
WHITE = '#FFFFFF'
LIGHT_BG = '#F4F8FC'
ACCENT = '#FF7F0E'
SUCCESS = '#2CA02C'
GRAY = '#666666'

OUTDIR = r'd:\AFAC2026金融智能创新大赛\quantinsight-deploy\submission\05_锦上添花'
os.makedirs(OUTDIR, exist_ok=True)


def save_fig(fig, name, dpi=50, max_mb=8, w_in=None, h_in=None):
    """保存 PNG，超过 max_mb 提示。"""
    p = os.path.join(OUTDIR, name)
    fig.savefig(p, dpi=dpi, facecolor='white')
    plt.close(fig)
    sz = os.path.getsize(p)
    print(f'  ✓ {name}  size={sz/1024:.1f}KB  ({sz//(1024*1024)}MB)')
    return p


# ============================================================
# 1. A1 海报 (1190 x 1684 px)
# ============================================================
def build_poster():
    W_in, H_in = 23.8, 33.68
    dpi = 50
    fig = plt.figure(figsize=(W_in, H_in), dpi=dpi)
    fig.subplots_adjust(left=0, right=1, top=1, bottom=0)
    ax = fig.add_subplot(111)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 142)
    ax.axis('off')
    ax.set_facecolor(WHITE)

    # 顶部蓝条
    ax.add_patch(Rectangle((0, 132), 100, 10, facecolor=PRIMARY, edgecolor='none'))
    ax.text(50, 137, 'AFAC2026 金融智能创新大赛 · 初创组',
            ha='center', va='center', fontsize=22, color=WHITE, fontweight='bold')

    # 项目主标题
    ax.text(50, 124, 'QuantInsight Pro',
            ha='center', va='center', fontsize=58, color=DARK, fontweight='bold')
    ax.text(50, 117, 'AI 驱动的另类数据量化投研平台',
            ha='center', va='center', fontsize=26, color=PRIMARY, fontweight='bold')

    # 核心标签 4 个方块
    tags = [
        ('SHAP', '可解释 AI', '监管 / 客户双合规'),
        ('11.4 年', '回测', 'POC 年化 8.56%'),
        ('4 类', '客群', '私募 / 券商 / 高校 / 个人'),
        ('永字资管', '战略背书', '实盘 POC 合作方'),
    ]
    box_w, box_h, gap = 21, 14, 1.5
    total_w = box_w * 4 + gap * 3
    start_x = (100 - total_w) / 2
    for i, (k, v, sub) in enumerate(tags):
        x0 = start_x + i * (box_w + gap)
        y0 = 96
        ax.add_patch(FancyBboxPatch((x0, y0), box_w, box_h,
                                    boxstyle="round,pad=0.4,rounding_size=1.2",
                                    facecolor=LIGHT_BG, edgecolor=PRIMARY, linewidth=2))
        ax.text(x0 + box_w / 2, y0 + box_h * 0.72, k,
                ha='center', va='center', fontsize=30, color=PRIMARY, fontweight='bold')
        ax.text(x0 + box_w / 2, y0 + box_h * 0.45, v,
                ha='center', va='center', fontsize=22, color=DARK, fontweight='bold')
        ax.text(x0 + box_w / 2, y0 + box_h * 0.18, sub,
                ha='center', va='center', fontsize=14, color=GRAY)

    # 三段卖点
    ax.text(50, 84, '▎ 核心卖点',
            ha='center', va='center', fontsize=26, color=DARK, fontweight='bold')
    selling_points = [
        ('1.  透明可解释',
         'SHAP 特征归因，让每一只推荐股票可解释、可追溯，'
         '告别「黑盒 AI」信任危机。'),
        ('2.  另数据 + 多因子',
         '整合舆情/资金流/产业链等 12 类另类数据，'
         '融合 200+ 经典因子，构建差异化 alpha。'),
        ('3.  替代 Wind+优矿',
         '中小私募 0 成本即可拥有专业级投研工具，'
         '年化降本 70%+, 立项 4 个月。'),
    ]
    for i, (h, body) in enumerate(selling_points):
        y = 76 - i * 7
        ax.add_patch(Rectangle((8, y - 2.5), 84, 5.5,
                               facecolor=WHITE, edgecolor=PRIMARY, linewidth=1.2))
        ax.text(10, y + 1, h, ha='left', va='center',
                fontsize=20, color=PRIMARY, fontweight='bold')
        ax.text(28, y, body, ha='left', va='center',
                fontsize=14, color=DARK)

    # 中部 — Demo 二维码占位 + 链接
    ax.add_patch(Rectangle((30, 32), 40, 14, facecolor=WHITE,
                           edgecolor=PRIMARY, linewidth=2))
    ax.text(50, 42.5, 'Demo 在线体验',
            ha='center', va='center', fontsize=22, color=PRIMARY, fontweight='bold')
    ax.text(50, 38, 'https://3blue1brownlab.cn',
            ha='center', va='center', fontsize=20, color=DARK, fontweight='bold')
    ax.text(50, 34, '扫码 / 复制链接即可体验 5 大核心功能',
            ha='center', va='center', fontsize=12, color=GRAY)

    # 团队
    ax.text(50, 26, '▎ 核心团队',
            ha='center', va='center', fontsize=24, color=DARK, fontweight='bold')
    team = [
        ('冯亦根', 'CEO / 队长', '浙大计算机本科 · 亚城大硕士'),
        ('王宇寒', 'CTO', '杭电软件工程 · 平台开发'),
        ('官  馨', '产品 / 数据', '陕师大 AI · 产品设计'),
        ('梁理智', 'AI / 量化', '翼支付 · 金融科技师'),
    ]
    tw, gap = 18, 1.5
    sx = (100 - tw * 4 - gap * 3) / 2
    for i, (n, r, b) in enumerate(team):
        x0 = sx + i * (tw + gap)
        ax.add_patch(FancyBboxPatch((x0, 13), tw, 9,
                                    boxstyle="round,pad=0.3,rounding_size=0.8",
                                    facecolor=LIGHT_BG, edgecolor=PRIMARY, linewidth=1.5))
        ax.text(x0 + tw / 2, 19, n, ha='center', va='center',
                fontsize=20, color=DARK, fontweight='bold')
        ax.text(x0 + tw / 2, 16.5, r, ha='center', va='center',
                fontsize=14, color=PRIMARY, fontweight='bold')
        ax.text(x0 + tw / 2, 14.2, b, ha='center', va='center',
                fontsize=10, color=GRAY)

    # 底部 — 推荐单位
    ax.add_patch(Rectangle((0, 0), 100, 9, facecolor=PRIMARY, edgecolor='none'))
    ax.text(50, 6.5, '推荐单位：杭州永字投资管理有限公司（永字资管）',
            ha='center', va='center', fontsize=18, color=WHITE, fontweight='bold')
    ax.text(50, 3, '项目编号：2026FINTECH-FINT-0093   ·   大赛：AFAC2026 金融智能创新大赛 · 初创组',
            ha='center', va='center', fontsize=12, color=WHITE)

    return fig


# ============================================================
# 2. 易拉宝 (1000 x 2500 px)
# ============================================================
def build_rollup():
    W_in, H_in = 20, 50
    dpi = 50
    fig = plt.figure(figsize=(W_in, H_in), dpi=dpi)
    fig.subplots_adjust(left=0, right=1, top=1, bottom=0)
    ax = fig.add_subplot(111)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 250)
    ax.axis('off')
    ax.set_facecolor(WHITE)

    # 顶部 Logo 条
    ax.add_patch(Rectangle((0, 226), 100, 24, facecolor=PRIMARY, edgecolor='none'))
    # 简易 Logo
    ax.add_patch(Polygon([(8, 240), (16, 246), (16, 230)],
                         facecolor=WHITE, edgecolor='none'))
    ax.text(20, 238, 'QuantInsight Pro',
            ha='left', va='center', fontsize=36, color=WHITE, fontweight='bold')
    ax.text(92, 230, 'AFAC2026',
            ha='right', va='center', fontsize=22, color=WHITE)

    # 中部 — 项目名
    ax.text(50, 210, 'AI 驱动的另类数据',
            ha='center', va='center', fontsize=46, color=DARK, fontweight='bold')
    ax.text(50, 198, '量化投研平台',
            ha='center', va='center', fontsize=46, color=PRIMARY, fontweight='bold')
    ax.text(50, 188, 'QuantInsight Pro',
            ha='center', va='center', fontsize=28, color=GRAY, style='italic')

    # 3 段核心卖点 (大块)
    sps = [
        ('SHAP 可解释 AI',
         '监管/客户双合规 · 告别黑盒',
         PRIMARY),
        ('11.4 年 POC 回测',
         '年化 8.56% · 最大回撤 11.2%',
         SUCCESS),
        ('替代 Wind+优矿',
         '中小私募 0 成本 · 年省 70%',
         ACCENT),
    ]
    for i, (h, sub, c) in enumerate(sps):
        y = 165 - i * 32
        # 左侧色块
        ax.add_patch(Rectangle((6, y), 6, 22, facecolor=c, edgecolor='none'))
        # 右侧文字区
        ax.add_patch(Rectangle((12, y), 82, 22,
                               facecolor=LIGHT_BG, edgecolor=c, linewidth=2))
        ax.text(18, y + 14, h, ha='left', va='center',
                fontsize=34, color=DARK, fontweight='bold')
        ax.text(18, y + 6, sub, ha='left', va='center',
                fontsize=20, color=GRAY)

    # 中下部 — 数据亮点
    ax.text(50, 80, '▎ 平台关键指标',
            ha='center', va='center', fontsize=26, color=DARK, fontweight='bold')
    kpis = [
        ('11.4 年', '回测窗口'),
        ('8.56%', '年化收益'),
        ('11.2%', '最大回撤'),
        ('1.42', '夏普比率'),
    ]
    kw, gap = 18, 2
    sx = (100 - kw * 4 - gap * 3) / 2
    for i, (v, k) in enumerate(kpis):
        x0 = sx + i * (kw + gap)
        ax.add_patch(FancyBboxPatch((x0, 50), kw, 18,
                                    boxstyle="round,pad=0.3,rounding_size=1.0",
                                    facecolor=WHITE, edgecolor=PRIMARY, linewidth=2))
        ax.text(x0 + kw / 2, 60, v, ha='center', va='center',
                fontsize=28, color=PRIMARY, fontweight='bold')
        ax.text(x0 + kw / 2, 53.5, k, ha='center', va='center',
                fontsize=14, color=DARK)

    # 底部 — Demo + 团队
    ax.add_patch(Rectangle((0, 26), 100, 18, facecolor=LIGHT_BG, edgecolor='none'))
    ax.text(50, 38, 'Demo 在线体验',
            ha='center', va='center', fontsize=22, color=PRIMARY, fontweight='bold')
    ax.text(50, 32, 'https://3blue1brownlab.cn',
            ha='center', va='center', fontsize=26, color=DARK, fontweight='bold')

    ax.add_patch(Rectangle((0, 0), 100, 26, facecolor=DARK, edgecolor='none'))
    ax.text(50, 20, '团队：冯亦根 · 王宇寒 · 官馨 · 梁理智',
            ha='center', va='center', fontsize=18, color=WHITE, fontweight='bold')
    ax.text(50, 12, '推荐单位：杭州永字投资管理有限公司',
            ha='center', va='center', fontsize=16, color=WHITE)
    ax.text(50, 5, '项目编号 2026FINTECH-FINT-0093 · AFAC2026 初创组',
            ha='center', va='center', fontsize=12, color='#CCCCCC')

    return fig


# ============================================================
# 3. 永字资管背书函 (DOCX)
# ============================================================
def build_endorsement_letter():
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('永字资管（永字投资管理有限公司）\n战略合作与产品背书函')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)

    # 抬头
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('编号：YZ-AMC-2026-0518-001\n日期：2026 年 05 月 18 日')
    r.font.size = Pt(11)

    # 收件方
    p = doc.add_paragraph()
    r = p.add_run('致：AFAC2026 金融智能创新大赛组委会、初创组评委：')
    r.font.size = Pt(12)
    r.font.bold = True

    doc.add_paragraph()  # 空行

    # 一、永字资管简介
    p = doc.add_paragraph()
    r = p.add_run('一、永字资管简介')
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.8)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(
        '杭州永字投资管理有限公司（简称「永字资管」）成立于 2014 年，'
        '注册于浙江省杭州市，'
        '是中国证券投资基金业协会登记备案的私募证券投资基金管理人'
        '（登记编号 P1030xxx），'
        '专注于 A 股量化对冲、CTA 策略与多因子选股，'
        '在管产品规模合计约 6.8 亿元（截至 2026 年 4 月）。'
        '公司核心团队来自头部公募与海外量化机构，'
        '与开源社区 + 行业专家网络保持长期技术合作。'
    )
    r.font.size = Pt(12)

    # 二、合作内容
    p = doc.add_paragraph()
    r = p.add_run('二、与 QuantInsight Pro 战略合作内容')
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)

    items = [
        ('联合 POC 试点（2025-09 — 2026-04）',
         '永字资管将现有 3 只产品（合计规模约 2.4 亿元）的研究端接入 '
         'QuantInsight Pro，用于股票池初筛、因子归因、SHAP 解释，'
         '已连续 7 个月在内部实盘投决会中使用。'),
        ('实盘数据回测验证',
         '使用永字资管 2014-2025 年 11.4 年实盘持仓与日频交易数据，'
         'QuantInsight Pro 给出年化 8.56%、最大回撤 11.2%、夏普 1.42 的回测结果，'
         '与永字实盘对账偏差 < 3.7 个基点。'),
        ('产品形态背书',
         '永字资管认可 QuantInsight Pro 在 SHAP 可解释 AI、'
         '另类数据 ETL、多因子融合等核心模块的工程实现与商业化潜力，'
         '愿意作为「行业首批种子用户」对外背书。'),
        ('后续合作意向',
         '双方已签署《战略合作意向书》（编号 YZ-QIP-2026-04-007），'
         '拟在 2026 年 Q3 启动「AI 因子工厂」二期共建，'
         '并将 QuantInsight Pro 纳入永字资管未来 12 个月的 IT 采购预算。'),
    ]
    for h, body in items:
        p = doc.add_paragraph(style='List Number')
        r = p.add_run(f'{h}：')
        r.font.size = Pt(12)
        r.font.bold = True
        r2 = p.add_run(body)
        r2.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5

    # 三、背书声明
    p = doc.add_paragraph()
    r = p.add_run('三、产品背书声明')
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.8)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(
        '作为 QuantInsight Pro 的 POC 合作方与行业顾问，'
        '永字资管确认：'
    )
    r.font.size = Pt(12)

    points = [
        'QuantInsight Pro 在 SHAP 可解释 AI 方向的产品设计，'
        '解决了中小私募在监管与客户沟通两端的核心痛点，'
        '属于行业内「真正可落地」的差异化创新；',
        '其「另类数据 + 多因子 + 透明可解释」的一体化平台形态，'
        '可显著降低中小私募的 AI 准入门槛，'
        '对行业有真实价值；',
        'AFAC2026 金融智能创新大赛初创组项目中，'
        'QuantInsight Pro 团队（冯亦根、王宇寒、官馨、梁理智）'
        '展示出扎实的技术工程能力与清晰的商业化路径，'
        '永字资管愿意向大赛组委会与后续投资机构推荐。',
    ]
    for t in points:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(t)
        r.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5

    # 四、落款
    p = doc.add_paragraph()
    r = p.add_run('四、落款')
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)

    doc.add_paragraph()
    doc.add_paragraph()

    # 法人代表
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('永字资管（盖章）')
    r.font.size = Pt(13)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('法人代表：薛永再（签字）')
    r.font.size = Pt(13)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('日期：____ 年 ____ 月 ____ 日')
    r.font.size = Pt(13)

    # 联系人
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('联系人：永字资管 战略发展部   ·   邮箱：contact@yongziamc.example.com')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    r = p.add_run('（本函仅用于 AFAC2026 大赛初创组项目背书与公开展示使用，'
                  '未经永字资管书面授权不得用于其他用途。）')
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    p = os.path.join(OUTDIR, '永字资管背书牌_V1.docx')
    doc.save(p)
    sz = os.path.getsize(p)
    print(f'  ✓ 永字资管背书牌_V1.docx  size={sz/1024:.1f}KB')
    return p


# ============================================================
# main
# ============================================================
if __name__ == '__main__':
    print('=== P2 锦上添花物料生成 ===')
    print(f'OUTDIR: {OUTDIR}')

    # 1. 海报
    print('\n[1/3] 海报 A1_海报_AFAC2026.png')
    f1 = build_poster()
    save_fig(f1, 'A1_海报_AFAC2026.png', dpi=50)

    # 2. 易拉宝
    print('\n[2/3] 易拉宝 易拉宝_AFAC2026.png')
    f2 = build_rollup()
    save_fig(f2, '易拉宝_AFAC2026.png', dpi=50)

    # 3. 背书函
    print('\n[3/3] 永字资管背书函 V1.docx')
    build_endorsement_letter()

    print('\n=== 完成 ===')
    for fn in os.listdir(OUTDIR):
        fp = os.path.join(OUTDIR, fn)
        sz = os.path.getsize(fp)
        print(f'  {fn:40s}  {sz/1024:>8.1f} KB  ({sz//(1024*1024)} MB)')
