# -*- coding: utf-8 -*-
"""Scan generated DOCX for leftover markdown artifacts."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
WORD_DIR = ROOT / "submission" / "03_正式文档_WORD"
README_DOCX = ROOT / "submission" / "00_项目README_QuantInsight_Pro.docx"
SELF_SCORE = ROOT / "submission" / "07_AFAC2026_自评打分报告.docx"

PATTERNS = {
    "hash_heading": re.compile(r"^\s*#{1,6}\s"),
    "bold_marker": re.compile(r"\*\*"),
    "backtick_fence": re.compile(r"```"),
    "md_link": re.compile(r"\[[^\]]+\]\([^)]+\)"),
    "table_pipe_row": re.compile(r"^\s*\|.+?\|\s*$"),
}

SKIP_FILES = set()  # e.g. legacy; empty = scan all


def _is_code_paragraph(paragraph) -> bool:
    runs = [r for r in paragraph.runs if r.text.strip()]
    if not runs:
        return False
    return all(r.font.name == "Consolas" for r in runs)


def scan_docx(path: Path) -> dict:
    doc = Document(path)
    counts = {k: 0 for k in PATTERNS}
    samples: list[str] = []

    def check_text(text: str, *, is_code: bool = False):
        if not text or is_code:
            return
        for key, pat in PATTERNS.items():
            if pat.search(text):
                counts[key] += len(pat.findall(text))
                if len(samples) < 5:
                    samples.append(text[:120])

    for p in doc.paragraphs:
        check_text(p.text, is_code=_is_code_paragraph(p))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                check_text(cell.text)
    counts["total"] = sum(counts.values())
    return {"path": path, "counts": counts, "samples": samples}


def main() -> int:
    targets = sorted(WORD_DIR.glob("*.docx"))
    if README_DOCX.exists():
        targets.append(README_DOCX)
    if SELF_SCORE.exists():
        targets.append(SELF_SCORE)

    print("=== DOCX Markdown Artifact Scan ===")
    failed = 0
    for path in targets:
        if path.name in SKIP_FILES:
            continue
        result = scan_docx(path)
        total = result["counts"]["total"]
        status = "PASS" if total == 0 else "FAIL"
        if total:
            failed += 1
        print(f"{status}  {path.name}: total={total}  {result['counts']}")
        for s in result["samples"][:2]:
            print(f"       sample: {s}")

    print(f"\nScanned {len(targets)} files · {failed} with artifacts")
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
