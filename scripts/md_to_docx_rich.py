# -*- coding: utf-8 -*-
"""Markdown → native Word rich text (no markdown artifacts)."""
from __future__ import annotations

import re
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DARK = RGBColor(0x0A, 0x0E, 0x27)
ACCENT = RGBColor(0x00, 0xD4, 0xFF)
MUTED = RGBColor(0x66, 0x66, 0x66)
LINK_COLOR = RGBColor(0x05, 0x63, 0xC1)
HEADER_BG = "0A0E27"
ROW_ALT = "E8F4FC"

INLINE_RE = re.compile(
    r"(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_|`.+?`|\[.+?\]\(.+?\))",
    re.DOTALL,
)
LINK_RE = re.compile(r"\[(.+?)\]\((.+?)\)")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
ORDERED_RE = re.compile(r"^(\d+)\.\s+(.+)$")
BULLET_RE = re.compile(r"^[-*+]\s+(.+)$")
TABLE_SEP_RE = re.compile(r"^\|?[\s\-:|]+\|?$")


def _font(run, *, size=10.5, bold=False, italic=False, color=None, name="微软雅黑", mono=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Consolas" if mono else name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑" if not mono else "Consolas")
    if color:
        run.font.color.rgb = color


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline_runs(paragraph, text: str, *, base_size=10.5):
    if not text:
        return
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos : m.start()])
            _font(r, size=base_size)
        chunk = m.group(0)
        if chunk.startswith("**") and chunk.endswith("**"):
            r = paragraph.add_run(chunk[2:-2])
            _font(r, size=base_size, bold=True)
        elif chunk.startswith("__") and chunk.endswith("__"):
            r = paragraph.add_run(chunk[2:-2])
            _font(r, size=base_size, bold=True)
        elif chunk.startswith("*") and chunk.endswith("*") and not chunk.startswith("**"):
            r = paragraph.add_run(chunk[1:-1])
            _font(r, size=base_size, italic=True)
        elif chunk.startswith("_") and chunk.endswith("_") and not chunk.startswith("__"):
            r = paragraph.add_run(chunk[1:-1])
            _font(r, size=base_size, italic=True)
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = paragraph.add_run(chunk[1:-1])
            _font(r, size=base_size - 0.5, mono=True, color=MUTED)
        elif chunk.startswith("["):
            lm = LINK_RE.match(chunk)
            if lm:
                add_hyperlink(paragraph, lm.group(1), lm.group(2))
            else:
                r = paragraph.add_run(chunk)
                _font(r, size=base_size)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        _font(r, size=base_size)


def shade(cell, color: str) -> None:
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    )


def cell_text(cell, text, *, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    add_inline_runs(p, str(text), base_size=size)
    for r in p.runs:
        if bold:
            r.bold = True
        if color:
            r.font.color.rgb = color


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    from docx.enum.table import WD_TABLE_ALIGNMENT

    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        shade(t.rows[0].cells[i], HEADER_BG)
        cell_text(
            t.rows[0].cells[i],
            h,
            bold=True,
            size=10,
            color=RGBColor(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ri % 2 == 0:
                shade(t.rows[ri + 1].cells[ci], ROW_ALT)
            cell_text(t.rows[ri + 1].cells[ci], val, bold=(ci == 0), size=9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def setup_doc(title: str, subtitle: str = "", meta: str = "") -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    _font(r, size=22, bold=True, color=DARK)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        _font(r2, size=12, color=ACCENT)
    if meta:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rm = mp.add_run(meta)
        _font(rm, size=9, color=MUTED)
    doc.add_paragraph()
    return doc


def heading(doc: Document, text: str, level: int = 1):
    clean = _strip_md_inline(text)
    h = doc.add_heading(clean, level=min(level, 3))
    for r in h.runs:
        _font(r, size=16 - level, bold=True, color=DARK if level == 1 else ACCENT)


def body_paragraph(doc: Document, text: str, *, quote=False, code=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(6)
    if quote:
        p.paragraph_format.left_indent = Cm(0.8)
        add_inline_runs(p, _strip_md_inline(text), base_size=10)
        for r in p.runs:
            r.italic = True
            r.font.color.rgb = MUTED
    elif code:
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(text)
        _font(r, size=9, mono=True, color=DARK)
    else:
        add_inline_runs(p, text, base_size=10.5)
    return p


def bullet(doc: Document, text: str, *, numbered=False, num_text=""):
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    if numbered and num_text:
        p.text = ""
    add_inline_runs(p, text, base_size=10.5)


def _strip_md_inline(text: str) -> str:
    text = LINK_RE.sub(r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


def _parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def _skip_frontmatter(lines: list[str]) -> list[str]:
    if not lines or lines[0].strip() != "---":
        return lines
    for i in range(1, len(lines)):
        if lines[i].strip() == "---":
            return lines[i + 1 :]
    return lines


def _unwrap_blockquote_prefix(line: str) -> tuple[str, bool]:
    """Return (content, was_blockquote)."""
    s = line.rstrip()
    if s.lstrip().startswith(">"):
        inner = re.sub(r"^\s*>\s?", "", s.lstrip(), count=1)
        return inner.strip(), True
    return s.strip(), False


def _is_table_row(line: str) -> bool:
    return line.strip().startswith("|")


def _is_table_separator(line: str) -> bool:
    return bool(TABLE_SEP_RE.match(line.strip()))


def convert_markdown_file(
    md_path: Path,
    out_path: Path,
    *,
    doc_title: str | None = None,
    subtitle: str = "",
    meta: str = "",
    on_section: Callable[[Document, str, int], None] | None = None,
) -> Path:
    """Convert markdown file to rich Word document."""
    lines = _skip_frontmatter(md_path.read_text(encoding="utf-8").splitlines())
    title = doc_title or md_path.stem
    doc = setup_doc(title, subtitle, meta)
    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                in_code = False
                if code_lines:
                    cap = doc.add_paragraph()
                    cap.paragraph_format.space_before = Pt(4)
                    label = f"代码块{f' ({code_lang})' if code_lang else ''}"
                    r = cap.add_run(label)
                    _font(r, size=9, bold=True, color=MUTED)
                    for cl in code_lines:
                        body_paragraph(doc, cl, code=True)
                code_lines = []
                code_lang = ""
            i += 1
            continue

        if in_code:
            code_lines.append(line.rstrip())
            i += 1
            continue

        if not stripped or stripped == "---":
            i += 1
            continue

        content, from_quote = _unwrap_blockquote_prefix(line)

        if _is_table_row(content) and i + 1 < len(lines):
            next_content, _ = _unwrap_blockquote_prefix(lines[i + 1])
            if _is_table_separator(next_content):
                headers = _parse_table_row(content)
                rows = []
                i += 2
                while i < len(lines):
                    row_content, _ = _unwrap_blockquote_prefix(lines[i])
                    if not _is_table_row(row_content):
                        break
                    rows.append(_parse_table_row(row_content))
                    i += 1
                if headers and rows:
                    add_table(doc, headers, rows)
                continue

        if from_quote and content:
            body_paragraph(doc, content, quote=True)
            i += 1
            continue

        hm = HEADING_RE.match(stripped)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2).strip()
            heading(doc, text, level=level)
            if on_section:
                on_section(doc, text, level)
            i += 1
            continue

        if stripped.startswith("<!--"):
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        if _is_table_row(stripped) and i + 1 < len(lines) and _is_table_separator(lines[i + 1].strip()):
            headers = _parse_table_row(stripped)
            rows = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_parse_table_row(lines[i]))
                i += 1
            if headers and rows:
                add_table(doc, headers, rows)
            continue

        if stripped.startswith(">"):
            body_paragraph(doc, stripped.lstrip("> ").strip(), quote=True)
            i += 1
            continue

        om = ORDERED_RE.match(stripped)
        if om:
            bullet(doc, om.group(2), numbered=True, num_text=om.group(1))
            i += 1
            continue

        bm = BULLET_RE.match(stripped)
        if bm:
            bullet(doc, bm.group(1))
            i += 1
            continue

        img = re.match(r"^!\[(.*?)\]\((.+?)\)", stripped)
        if img:
            alt, src = img.group(1), img.group(2)
            img_path = Path(src)
            if not img_path.is_absolute():
                img_path = md_path.parent / src
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(img_path), width=Cm(14))
                if alt:
                    cap = doc.add_paragraph(alt)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in cap.runs:
                        _font(r, size=9, color=MUTED)
            i += 1
            continue

        if stripped.startswith("!["):
            i += 1
            continue

        body_paragraph(doc, stripped)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
