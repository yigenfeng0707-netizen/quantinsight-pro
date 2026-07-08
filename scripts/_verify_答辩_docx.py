# -*- coding: utf-8 -*-
"""Verify the 5 generated DOCX files."""
from pathlib import Path

from docx import Document

BASE = Path(r"d:\AFAC2026金融智能创新大赛\quantinsight-deploy\submission\03_正式文档_WORD")
FILES = [
    "09_答辩话术_V3.docx",
    "10_3轮模拟答辩_V1.docx",
    "11_风险预案_V2.docx",
    "12_5杀手锏提问_V1.docx",
    "13_QA_Database_V2.docx",
]

print(f"{'文件':<28} {'段落':>6} {'表格':>6} {'大小(KB)':>10}")
print("-" * 60)
for n in FILES:
    p = BASE / n
    if not p.exists():
        print(f"{n:<28} MISSING")
        continue
    d = Document(p)
    size_kb = round(p.stat().st_size / 1024, 1)
    print(f"{n:<28} {len(d.paragraphs):>6} {len(d.tables):>6} {size_kb:>10}")
