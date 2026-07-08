# -*- coding: utf-8 -*-
"""生成 submission/00_项目README_QuantInsight_Pro.docx 与 .html"""
from __future__ import annotations

import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SUB = ROOT / "submission"
POC_JSON = SUB / "02_Demo交付" / "POC实验数据" / "t35_hs300_summary.json"
DOCX_OUT = SUB / "00_项目README_QuantInsight_Pro.docx"
HTML_OUT = SUB / "00_项目README_QuantInsight_Pro.html"
CHART_DIR = SUB / "03_正式文档_WORD" / "_charts"

DARK = RGBColor(0x0A, 0x0E, 0x27)
ACCENT = RGBColor(0x00, 0xD4, 0xFF)
GOLD = RGBColor(0xFF, 0xB8, 0x00)
HEADER_BG = "0A0E27"
ROW_ALT = "E8F4FC"
PROJECT_ID = "2026FINTECH-FINT-0093"
PLATFORM_ID = "20260110040"
DEMO_URL = "https://3blue1brownlab.cn"


def shade(cell, color: str) -> None:
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    )


def cell_text(cell, text, *, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if color:
        r.font.color.rgb = color


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        shade(t.rows[0].cells[i], HEADER_BG)
        cell_text(
            t.rows[0].cells[i], h, bold=True, size=10,
            color=RGBColor(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ri % 2 == 0:
                shade(t.rows[ri + 1].cells[ci], ROW_ALT)
            cell_text(t.rows[ri + 1].cells[ci], val, bold=(ci == 0), size=9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r.font.color.rgb = DARK if level == 1 else ACCENT


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.size = Pt(10.5)
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_picture(doc, img_path: Path, caption: str = "", width_cm: float = 15.0):
    if not img_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            r.font.name = "微软雅黑"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    doc.add_paragraph()


def load_poc():
    with open(POC_JSON, encoding="utf-8") as f:
        return json.load(f)


def team_rows():
    return [
        [
            "CEO / 主讲",
            "冯亦根",
            "浙江大学计算机通信工程本科、亚洲城市大学硕士；慧点资本创始人，CFA；15+ 年金融科技投资经验。",
            "战略、融资、BD、路演总负责",
        ],
        [
            "CTO",
            "王宇寒",
            "杭州电子科技大学软件工程专业 2022 级本科；平台架构与 AI 工程负责人。",
            "技术架构、AI 模型、系统开发、生产部署",
        ],
        [
            "产品 / 数据",
            "官馨",
            "陕西师范大学人工智能专业大三；产品设计与客户研究。",
            "数据策略、产品设计、客户研究、UX",
        ],
        [
            "AI / 量化",
            "梁理智",
            "河北科技大学；翼支付 AI 开发者，金融科技师（二级）证书。",
            "量化策略、AI 金融科技落地、运营",
        ],
    ]


def module_rows():
    return [
        ["1", "智能选股", "17 因子综合评分 + Top10 推荐", "✅ 已上线"],
        ["2", "SHAP 解读", "4 类 SHAP 图 + 17 因子归因", "✅ 业内独家"],
        ["3", "AI 投研问答", "RAG 数据接地 + Qwen 解读 + 引用溯源", "✅ 已上线"],
        ["4", "量化策略回测", "5 策略 + 11.4 年 HS300 真实数据", "✅ 开源引擎"],
        ["5", "智能盯盘", "7×24 涨跌停预警 + 异动监控", "✅ 已上线"],
        ["6", "模拟交易", "A 股实时模拟 + 3 层 fallback", "✅ 已上线"],
        ["7", "自动报告", "6 段式周报 + Word/PDF 导出", "✅ 已上线"],
        ["8", "实时数据看板", "大盘卡片 + 北向资金热力图", "✅ V2 已上线"],
    ]


def submission_dir_rows():
    return [
        ["01_商业计划书_QuantInsight_Pro.md/.html", "AFAC 八章节商业计划书（权威）"],
        ["02_Demo交付/", "Demo 视频、运行指南、POC 数据、交互设计"],
        ["03_正式文档_WORD/", "9 份正式 Word 文档（规则对照、测试、FAQ 等）"],
        ["04_测试报告/", "单元测试与冒烟测试 JSON"],
        ["05_团队信息一致性核查报告.md", "团队口径审计报告"],
        ["06_文档归档说明.md", "V1 文档归档与 README 位置说明"],
        ["00_项目README_QuantInsight_Pro.docx/.html", "本文件 — 项目一页纸总览"],
        ["QuantInsight_Pro_AFAC2026_提交包_*.zip", "一键打包提交 ZIP"],
    ]


def render_charts():
    sys.path.insert(0, str(ROOT / "scripts"))
    from afac_charts import chart_strategy_comparison, chart_afac_scoring_radar, chart_team_org

    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    CHART_DIR.mkdir(parents=True, exist_ok=True)
    strategy = chart_strategy_comparison("readme_strategy.png")
    scoring = chart_afac_scoring_radar("readme_scoring.png")
    team_org = chart_team_org("readme_team_org.png")

    poc = load_poc()
    mf = poc["strategies"]["multi_factor"]
    bh = poc["strategies"]["buy_hold"]
    labels = ["多因子年化", "基准年化", "超额收益", "夏普(多因子)"]
    values = [mf["annual_return_pct"], bh["annual_return_pct"], mf["excess_return_vs_benchmark_pct"], mf["sharpe"]]
    colors = ["#00D4FF", "#7B61FF", "#FFB800", "#00FF88"]
    fig, ax = plt.subplots(figsize=(8, 4.5), facecolor="white")
    bars = ax.bar(labels, values, color=colors, edgecolor="#0A0E27", linewidth=0.6)
    ax.set_title("T35 回测核心指标（HS300 · 2015–2026）", fontsize=12, fontweight="bold", color="#0A0E27")
    ax.set_ylabel("数值")
    ax.grid(axis="y", alpha=0.25, linestyle="--")
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05, f"{val}", ha="center", fontsize=9)
    fig.tight_layout()
    metrics_path = CHART_DIR / "readme_key_metrics.png"
    fig.savefig(metrics_path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return {"strategy": strategy, "scoring": scoring, "metrics": metrics_path, "team_org": team_org}


def zip_path_hint() -> str:
    today = datetime.now().strftime("%Y%m%d")
    p = SUB / f"QuantInsight_Pro_AFAC2026_提交包_{today}.zip"
    if p.exists():
        return str(p.relative_to(ROOT))
    zips = sorted(SUB.glob("QuantInsight_Pro_AFAC2026_提交包_*.zip"))
    return str(zips[-1].relative_to(ROOT)) if zips else f"submission/QuantInsight_Pro_AFAC2026_提交包_{today}.zip"


def generate_docx(charts: dict) -> Path:
    poc = load_poc()
    mf = poc["strategies"]["multi_factor"]
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("QuantInsight Pro")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = DARK
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("AI 驱动的另类数据量化投研平台 · AFAC2026 金融智能创新大赛 · 初创组")
    rs.font.size = Pt(12)
    rs.font.color.rgb = ACCENT
    rs.font.name = "微软雅黑"
    rs._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(
        f"项目编号 {PROJECT_ID} · AFAC 平台 ID {PLATFORM_ID} · {datetime.now():%Y年%m月%d日}"
    )
    rm.font.size = Pt(9)
    rm.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    heading(doc, "执行摘要")
    body(
        doc,
        "QuantInsight Pro 是面向专业机构投资者的新一代资管科技平台，业内首家将 SHAP 可解释性"
        "深度集成到 A 股智能选股流程。产品整合另类数据、AI 投研智能体与 MIT 开源回测引擎，"
        "为 9 类机构客户提供数据 + 模型 + 策略 + 合规的一站式智能投研解决方案。"
        "推荐单位杭州永字资产管理有限公司战略合作已签署，生产 Demo 已部署公网 HTTPS。"
    )

    heading(doc, "参赛团队（4 人 · AFAC 平台注册）")
    add_table(
        doc,
        ["角色", "姓名", "背景摘要", "职责"],
        team_rows(),
        [2.5, 2, 7, 4],
    )

    heading(doc, "核心功能模块（8 大模块）")
    add_table(doc, ["#", "模块", "能力", "Demo"], module_rows(), [1, 2.5, 7, 2.5])

    heading(doc, "关键指标")
    add_table(
        doc,
        ["指标", "数值", "说明"],
        [
            ["T35 多因子年化收益", f"{mf['annual_return_pct']}%", "HS300 · 11.4 年 akshare 真实数据"],
            ["夏普比率", str(mf["sharpe"]), "相对买入持有超额 3.10%"],
            ["最大回撤", f"{mf['max_drawdown_pct']}%", "较基准改善 33.97%"],
            ["单元测试", "21/21 PASS", "回测引擎 7/7 + 数据管道 14/14"],
            ["永字资管合作", "战略合作已签署", "推荐单位法定代表人薛永再 · 场外顾问"],
            ["AFAC 自评分", "88 / 100", "初筛五维加权自评"],
        ],
        [4, 3, 9],
    )
    add_picture(doc, charts["metrics"], "图1 T35 回测核心指标", width_cm=14)
    add_picture(doc, charts["strategy"], "图2 HS300 五策略回测对比（T35 修正）", width_cm=15)
    add_picture(doc, charts["team_org"], "图3 参赛团队组织架构", width_cm=14)

    heading(doc, "访问与提交")
    add_table(
        doc,
        ["项目", "路径 / URL"],
        [
            ["在线 Demo", DEMO_URL],
            ["GitHub", "github.com/yigenfeng0707-netizen/quantinsight-pro"],
            ["AFAC 控制台", f"https://afac.alipay.com/console/projects/{PLATFORM_ID}"],
            ["提交 ZIP", zip_path_hint()],
            ["文档归档", "archive/legacy_v1/（V1 过时材料，非提交依据）"],
        ],
        [4, 12],
    )

    heading(doc, "submission/ 目录指南")
    add_table(doc, ["路径", "说明"], submission_dir_rows(), [6, 10])

    heading(doc, "品牌色规范")
    add_table(
        doc,
        ["色名", "Hex", "用途"],
        [
            ["深空蓝", "#0A0E27", "主背景 / 标题 / 表格表头"],
            ["科技青", "#00D4FF", "强调色 / 链接 / 图表主色"],
            ["琥珀金", "#FFB800", "高亮 / 评分 / 次要图表"],
        ],
        [3, 3, 10],
    )
    body(doc, "联系人：冯亦根 ceo@3blue1brownlab.cn · 王宇寒 cto@3blue1brownlab.cn")

    SUB.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    print(f"  OK {DOCX_OUT.name} ({DOCX_OUT.stat().st_size:,} bytes)")
    return DOCX_OUT


def generate_html(charts: dict) -> Path:
    poc = load_poc()
    mf = poc["strategies"]["multi_factor"]
    zip_hint = zip_path_hint().replace("\\", "/")

    def img_tag(name: str, caption: str) -> str:
        rel = f"03_正式文档_WORD/_charts/{name}"
        return f'<figure><img src="{rel}" alt="{caption}"><figcaption>{caption}</figcaption></figure>'

    team_html = "".join(
        f"<tr><td>{r[0]}</td><td><strong>{r[1]}</strong></td><td>{r[2]}</td><td>{r[3]}</td></tr>"
        for r in team_rows()
    )
    mod_html = "".join(
        f"<tr><td>{r[0]}</td><td>{r[1]}</td><td>{r[2]}</td><td>{r[3]}</td></tr>"
        for r in module_rows()
    )
    dir_html = "".join(f"<tr><td><code>{r[0]}</code></td><td>{r[1]}</td></tr>" for r in submission_dir_rows())

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>QuantInsight Pro — 项目 README · AFAC2026</title>
<style>
:root {{ --dark:#0A0E27; --cyan:#00D4FF; --gold:#FFB800; --bg:#f4f8fc; }}
* {{ box-sizing:border-box; }}
body {{ font-family:"Microsoft YaHei","PingFang SC",sans-serif; background:var(--bg); color:#222; line-height:1.7; margin:0; }}
.wrap {{ max-width:920px; margin:0 auto; padding:32px 24px 48px; }}
header {{ background:var(--dark); color:#fff; padding:36px 32px; border-radius:12px; margin-bottom:28px; }}
header h1 {{ margin:0 0 8px; font-size:2rem; color:var(--cyan); }}
header p {{ margin:4px 0; opacity:.92; }}
.badge {{ display:inline-block; background:var(--gold); color:var(--dark); padding:2px 10px; border-radius:4px; font-size:.85rem; font-weight:700; margin-top:8px; }}
section {{ background:#fff; border-radius:10px; padding:24px 28px; margin-bottom:20px; box-shadow:0 2px 8px rgba(10,14,39,.06); }}
h2 {{ color:var(--dark); border-left:4px solid var(--cyan); padding-left:12px; margin-top:0; }}
table {{ width:100%; border-collapse:collapse; font-size:.92rem; margin:12px 0; }}
th {{ background:var(--dark); color:#fff; padding:10px 8px; text-align:left; }}
td {{ padding:9px 8px; border-bottom:1px solid #e8eef5; vertical-align:top; }}
tr:nth-child(even) td {{ background:#E8F4FC; }}
figure {{ text-align:center; margin:20px 0; }}
figure img {{ max-width:100%; border-radius:8px; border:1px solid #dde; }}
figcaption {{ font-size:.85rem; color:#666; margin-top:6px; }}
.colors {{ display:flex; gap:16px; flex-wrap:wrap; }}
.swatch {{ flex:1; min-width:140px; padding:16px; border-radius:8px; color:#fff; text-align:center; font-weight:600; }}
.swatch small {{ display:block; font-weight:400; opacity:.9; margin-top:4px; }}
footer {{ text-align:center; color:#666; font-size:.85rem; margin-top:24px; }}
@media print {{ body {{ background:#fff; }} section {{ box-shadow:none; border:1px solid #ddd; page-break-inside:avoid; }} }}
</style>
</head>
<body>
<div class="wrap">
<header>
  <h1>QuantInsight Pro</h1>
  <p>AI 驱动的另类数据量化投研平台</p>
  <p>AFAC2026 金融智能创新大赛 · 初创组 · 项目编号 {PROJECT_ID}</p>
  <p>AFAC 平台 ID {PLATFORM_ID} · Demo <a href="{DEMO_URL}" style="color:var(--cyan)">{DEMO_URL}</a></p>
  <span class="badge">永字资管战略合作已签署</span>
</header>

<section>
<h2>执行摘要</h2>
<p>QuantInsight Pro 面向专业机构投资者，将 SHAP 可解释性深度集成到 A 股智能选股流程，整合另类数据、AI 投研智能体与 MIT 开源回测引擎，提供数据 + 模型 + 策略 + 合规的一站式智能投研解决方案。生产 Demo 已公网部署，推荐单位杭州永字资产管理有限公司战略合作已签署。</p>
</section>

<section>
<h2>参赛团队（4 人）</h2>
<table><thead><tr><th>角色</th><th>姓名</th><th>背景</th><th>职责</th></tr></thead><tbody>{team_html}</tbody></table>
<p><em>推荐单位顾问：薛永再（杭州永字资管法定代表人，场外非参赛队员）</em></p>
</section>

<section>
<h2>核心功能（8 模块）</h2>
<table><thead><tr><th>#</th><th>模块</th><th>能力</th><th>状态</th></tr></thead><tbody>{mod_html}</tbody></table>
</section>

<section>
<h2>关键指标（T35 回测）</h2>
<table>
<thead><tr><th>指标</th><th>数值</th></tr></thead>
<tbody>
<tr><td>多因子年化收益</td><td><strong>{mf['annual_return_pct']}%</strong></td></tr>
<tr><td>夏普比率</td><td>{mf['sharpe']}</td></tr>
<tr><td>最大回撤</td><td>{mf['max_drawdown_pct']}%</td></tr>
<tr><td>单元测试</td><td>21/21 PASS</td></tr>
<tr><td>永字资管</td><td>战略合作已签署</td></tr>
<tr><td>AFAC 自评</td><td>88 / 100</td></tr>
</tbody>
</table>
{img_tag("readme_key_metrics.png", "图1 T35 回测核心指标")}
{img_tag("readme_strategy.png", "图2 HS300 五策略回测对比")}
</section>

<section>
<h2>访问与提交</h2>
<table>
<tbody>
<tr><td>在线 Demo</td><td><a href="{DEMO_URL}">{DEMO_URL}</a></td></tr>
<tr><td>提交 ZIP</td><td><code>{zip_hint}</code></td></tr>
<tr><td>文档归档</td><td><code>archive/legacy_v1/</code>（V1 过时材料）</td></tr>
<tr><td>权威 BP</td><td><code>submission/01_商业计划书_QuantInsight_Pro.md</code></td></tr>
</tbody>
</table>
</section>

<section>
<h2>submission/ 目录指南</h2>
<table><thead><tr><th>路径</th><th>说明</th></tr></thead><tbody>{dir_html}</tbody></table>
</section>

<section>
<h2>品牌色</h2>
<div class="colors">
  <div class="swatch" style="background:#0A0E27">深空蓝<small>#0A0E27 · 主背景/标题</small></div>
  <div class="swatch" style="background:#00D4FF;color:#0A0E27">科技青<small>#00D4FF · 强调/链接</small></div>
  <div class="swatch" style="background:#FFB800;color:#0A0E27">琥珀金<small>#FFB800 · 高亮/评分</small></div>
</div>
</section>

<footer>编制日期 {datetime.now():%Y年%m月%d日} · 冯亦根 ceo@3blue1brownlab.cn · 王宇寒 cto@3blue1brownlab.cn</footer>
</div>
</body>
</html>"""
    HTML_OUT.write_text(html, encoding="utf-8")
    print(f"  OK {HTML_OUT.name} ({HTML_OUT.stat().st_size:,} bytes)")
    return HTML_OUT


def main():
    print("=== 项目 README（Word + HTML）===")
    print("\n[1/2] 渲染图表...")
    charts = render_charts()
    print("\n[2/2] 生成文档...")
    generate_docx(charts)
    generate_html(charts)
    print(f"\n完成 → {DOCX_OUT.parent}")


if __name__ == "__main__":
    main()
